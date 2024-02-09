# EXPLAINATION: this is a program that summarize news headlines using a webscraper and then add the headlines
# to a document then saves it and automatically open it.

# Pseudocode

# import needed libraries
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Provide the link for the news website that would be webscraped
url = 'https://www.philstar.com/'

# get request
response = requests.get(url)
# Create an empty list for the website headlines so it will serve as a holder for the headlines
headlines_data = []
# parse html
if response.status_code == 200:
    soup = BeautifulSoup(response.text, 'html.parser')

    headline_tags = soup.find_all('h2', limit=4) 

# Extract and print the headlines
    for i, tag in enumerate(headline_tags, start=1):
        headline_link = tag.find('a') # the Philstar headline is located in the <a> tag so i used "a" to find it.
        if headline_link:
            article_url = headline_link.get('href')
            # Ensure the URL is absolute
            if article_url.startswith('/'):
                article_url = 'https://www.philstar.com' + article_url
            
# new modification: Get the 1st paragraph of the headline as a preview of the story.

            # Get the article page
            article_response = requests.get(article_url)
            if article_response.status_code == 200:
                article_soup = BeautifulSoup(article_response.text, 'html.parser')
                first_paragraph = article_soup.find('p')  # The Philstar 1st paragraph is located in the <p> tag so i used "p" to find it.
                if first_paragraph:
                    first_paragraph_text = first_paragraph.text.strip()
                else:
                    first_paragraph_text = "The preview for this article is currently unavailable."

                headlines_data.append({
                    'title': headline_link.text.strip(),
                    'url': article_url,
                    'first_paragraph': first_paragraph_text
                })

# Create a new Document
doc = Document()
doc.add_heading('Scraped Headlines', 0)

for item in headlines_data:
    # Add headline title
    title = doc.add_paragraph()
    run_title = title.add_run(item['title'])
    run_title.bold = True
    run_title.font.size = Pt(14)

    # Add URL 
    doc.add_paragraph(item['url'], style='IntenseQuote')

    # Add first paragraph
    first_paragraph = doc.add_paragraph(item['first_paragraph'])
    first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    first_paragraph.paragraph_format.space_after = Pt(18)  

    # Adding a line space for separation between articles
    doc.add_paragraph()

# Save the document
doc_filename = 'scraped_headlines.docx'
doc.save(doc_filename)

print("The headlines have been successfully scraped and saved into a Word document.")

# Automatically open the document on Windows
os.startfile(doc_filename)

# Note: in the future, add a function that would automatically print the document and also add a fuction
# that creates a barcode for every headline that if scanned would re-direct you to that specific headline